"""
Admin ERP RAG service — hybrid retrieval (BM25 + dense vectors) over uploaded data.

Separate from the policy knowledge base in knowledge_base.py. Only used when an
admin enables RAG mode on the admin dashboard.
"""

from __future__ import annotations

import io
import json
import logging
import os
import pickle
import re
import shutil
import threading
import time
import traceback
from typing import Any, Literal

import pandas as pd
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from rank_bm25 import BM25Okapi

logger = logging.getLogger("taia.rag_service")

_BASE_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "data", "admin_rag_store")
ADMIN_RAG_PERSIST_DIR = os.getenv("ADMIN_RAG_PERSIST_DIR", _BASE_DIR)
PERSIST_DIR = ADMIN_RAG_PERSIST_DIR
BM25_INDEX_PATH = os.getenv("ADMIN_BM25_INDEX_PATH", os.path.join(PERSIST_DIR, "bm25_index.pkl"))
MAX_UPLOAD_BYTES = int(os.getenv("ADMIN_RAG_MAX_UPLOAD_BYTES", str(10 * 1024 * 1024)))
RRF_K = 60
DENSE_CANDIDATES = 20
SPARSE_CANDIDATES = 20
_FAISS_INDEX_FILE = "index.faiss"

_lock = threading.Lock()
_embeddings = None
_vector_store: FAISS | None = None
_documents: list[Document] = []
_bm25: BM25Okapi | None = None


def _tokenize(text: str) -> list[str]:
    return re.findall(r"[a-z0-9]+", (text or "").lower())


def _create_embeddings():
    try:
        from langchain_huggingface import HuggingFaceEmbeddings
    except ImportError:
        from langchain_community.embeddings import HuggingFaceEmbeddings

    return HuggingFaceEmbeddings(
        model_name="all-MiniLM-L6-v2",
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )


def _get_embeddings():
    global _embeddings
    if _embeddings is None:
        _embeddings = _create_embeddings()
    return _embeddings


def _faiss_index_exists() -> bool:
    return os.path.isfile(os.path.join(PERSIST_DIR, _FAISS_INDEX_FILE))


def _load_faiss_index() -> FAISS:
    return FAISS.load_local(
        PERSIST_DIR,
        _get_embeddings(),
        allow_dangerous_deserialization=True,
    )


def _serialize_documents(docs: list[Document]) -> list[dict[str, Any]]:
    return [{"page_content": doc.page_content, "metadata": dict(doc.metadata)} for doc in docs]


def _deserialize_documents(payload: list[Any]) -> list[Document]:
    documents: list[Document] = []
    for item in payload:
        if isinstance(item, Document):
            documents.append(item)
        elif isinstance(item, dict):
            documents.append(
                Document(
                    page_content=str(item.get("page_content", "")),
                    metadata=dict(item.get("metadata") or {}),
                )
            )
        else:
            documents.append(Document(page_content=str(item), metadata={}))
    return documents


def _load_bm25_from_disk() -> None:
    """Restore BM25 index and document list from disk if present."""
    global _documents, _bm25
    if not os.path.exists(BM25_INDEX_PATH):
        return
    try:
        with open(BM25_INDEX_PATH, "rb") as handle:
            payload = pickle.load(handle)
        docs = payload.get("documents", [])
        if not docs:
            return
        _documents = _deserialize_documents(docs)
        tokenized = [_tokenize(doc.page_content) for doc in _documents]
        _bm25 = BM25Okapi(tokenized)
        logger.info("Admin RAG BM25 index loaded (%d documents)", len(_documents))
    except Exception as exc:
        logger.warning("Failed to load BM25 index: %s", exc)


def _save_bm25_to_disk() -> None:
    os.makedirs(os.path.dirname(BM25_INDEX_PATH) or PERSIST_DIR, exist_ok=True)
    with open(BM25_INDEX_PATH, "wb") as handle:
        pickle.dump({"documents": _serialize_documents(_documents)}, handle)


def _format_chunk(doc: Document) -> str:
    pages = doc.metadata.get("page")
    if not pages:
        return doc.page_content
    if isinstance(pages, int):
        pages = [pages]
    if len(pages) == 1:
        return f"[Page {pages[0]}] {doc.page_content}"
    return f"[Pages {', '.join(str(p) for p in pages)}] {doc.page_content}"


def _get_admin_vector_store() -> FAISS | None:
    global _vector_store
    if _vector_store is not None:
        return _vector_store
    if not os.path.isdir(PERSIST_DIR):
        return None
    try:
        if _faiss_index_exists():
            _vector_store = _load_faiss_index()
        else:
            os.makedirs(PERSIST_DIR, exist_ok=True)
            _vector_store = FAISS.from_texts(["init"], _get_embeddings())
            _vector_store.save_local(PERSIST_DIR)
        return _vector_store
    except Exception as exc:
        logger.warning("Admin RAG vector store unavailable: %s", exc)
        return None


def _row_to_document(row: Any) -> str | None:
    if isinstance(row, dict):
        parts = []
        for key, value in row.items():
            if value is None:
                continue
            text = str(value).strip()
            if not text or text.lower() == "nan":
                continue
            label = str(key).replace("_", " ").strip().title()
            parts.append(f"{label}: {text}")
        return ", ".join(parts) if parts else None
    text = str(row).strip()
    return text or None


def _clean_text(text: str) -> str:
    """Collapse excessive whitespace in extracted document text."""
    return re.sub(r"\s+", " ", (text or "")).strip()


# ── Natural boundary splitting with real offsets ──
_BULLET_PATTERN = re.compile(r"^\s*(?:•|-|\d+\.)\s+", re.MULTILINE)


def _split_into_natural_units(text: str) -> list[tuple[str, int, int]]:
    """
    Split text into natural units (paragraphs, bullet items).
    Returns list of (unit_text, start_char, end_char) where start_char is inclusive
    and end_char is exclusive in the original text.

    IMPORTANT: When a paragraph contains a bullet list, any heading/intro text that
    appears BEFORE the first bullet marker (e.g. "Civil Engineering Department
    Laboratories:") is prepended to every individual bullet item's text. This keeps
    section/department context attached to each list item even when that item is
    retrieved in isolation as its own chunk — without this, isolated bullet items
    like "Structures Lab" would carry no signal connecting them to "Civil
    Engineering", making them nearly unretrievable for department-scoped queries.
    Only the page_content carries the prepended heading; character offsets used for
    page-number metadata stay anchored to the bullet item's own true position.
    """
    units: list[tuple[str, int, int]] = []

    # Find all paragraphs by splitting on blank lines, with real offsets
    for match in re.finditer(r"(?s)(?:^|\n\s*\n)\s*([^\n].*?)(?=\n\s*\n|\Z)", text):
        para = match.group(1).strip()
        if not para:
            continue
        # para_start is the start of the non-whitespace content after the blank line
        para_start = match.start(1)

        # Check if this paragraph contains bullet lines
        bullet_lines = [m for m in _BULLET_PATTERN.finditer(para)]

        if len(bullet_lines) >= 1:
            # Text before the first bullet marker = heading/intro context for the list
            heading_text = para[: bullet_lines[0].start()].strip()

            # Split the paragraph into individual bullet items
            item_boundaries = []
            for i, m in enumerate(bullet_lines):
                start = m.start()
                if i + 1 < len(bullet_lines):
                    end = bullet_lines[i + 1].start()
                else:
                    end = len(para)
                item_boundaries.append((start, end))

            for local_start, local_end in item_boundaries:
                item_text = para[local_start:local_end].strip()
                if not item_text:
                    continue
                # Prepend heading context to the item's indexed text (content only —
                # offsets below stay tied to the bullet item itself, not the heading,
                # so page-number mapping remains accurate).
                if heading_text:
                    item_text = f"{heading_text} {item_text}"
                abs_start = para_start + local_start
                abs_end = para_start + local_end
                units.append((item_text, abs_start, abs_end))
        else:
            # No bullets, keep as whole paragraph
            para_end = match.end(1)
            units.append((para, para_start, para_end))

    return units


def _split_oversized_unit(unit_text: str, start_char: int, chunk_size: int) -> list[tuple[str, int, int]]:
    """
    Split a single unit that exceeds chunk_size using a plain sliding window,
    preserving real character offsets.
    """
    result = []
    unit_len = len(unit_text)
    overlap = 200
    pos = 0
    while pos < unit_len:
        chunk = unit_text[pos : pos + chunk_size].strip()
        if chunk:
            result.append((chunk, start_char + pos, start_char + pos + len(chunk)))
        pos += chunk_size - overlap
        if pos >= unit_len:
            break
    return result


def _parse_pdf_documents(file_bytes: bytes, filename: str) -> list[Document]:
    """
    Extract text from PDF using pdfplumber (better reading order, table handling).
    Chunks are created across page boundaries with overlap and consistent metadata.
    """
    try:
        import pdfplumber
    except ImportError:
        logger.error("pdfplumber not installed. Install with: pip install pdfplumber")
        return []

    documents: list[Document] = []
    full_text_parts: list[tuple[str, int]] = []  # (text, page_number)
    table_chunks: list[Document] = []

    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                # 1. Extract tables
                tables = page.extract_tables()
                for table in tables:
                    if not table or len(table) < 2:
                        continue
                    headers = [str(h).strip() if h else f"Col{i}" for i, h in enumerate(table[0])]
                    for row in table[1:]:
                        row_dict = {
                            headers[i]: (str(row[i]).strip() if i < len(row) and row[i] else "")
                            for i in range(len(headers))
                        }
                        text = _row_to_document(row_dict)
                        if text:
                            table_chunks.append(
                                Document(
                                    page_content=text,
                                    metadata={"page": [page_number], "source": "table"}
                                )
                            )

                # 2. Extract regular text
                text = (page.extract_text() or "").strip()
                if text:
                    full_text_parts.append((text, page_number))
    except Exception as exc:
        logger.error("pdfplumber failed to parse PDF %s: %s", filename, exc, exc_info=True)
        return []

    # Concatenate all page texts and build page offset ranges
    combined_text = ""
    page_offsets: list[tuple[int, int, int]] = []  # (start_char, end_char, page_number)
    for text, page_num in full_text_parts:
        start = len(combined_text)
        combined_text += text + "\n\n"
        end = len(combined_text)
        page_offsets.append((start, end, page_num))

    if not combined_text.strip():
        return table_chunks

    # Split into natural units with real character offsets
    raw_units = _split_into_natural_units(combined_text)

    chunk_size = 1000
    overlap_chars = 200

    # Split any unit that exceeds chunk_size
    natural_units: list[tuple[str, int, int]] = []
    for unit_text, unit_start, unit_end in raw_units:
        if len(unit_text) > chunk_size:
            natural_units.extend(_split_oversized_unit(unit_text, unit_start, chunk_size))
        else:
            natural_units.append((unit_text, unit_start, unit_end))

    # Build chunks with overlap
    current_chunk_units: list[tuple[str, int, int]] = []
    current_len = 0
    all_chunks: list[tuple[str, int, int]] = []  # (chunk_text, start_char, end_char)

    def _get_pages(char_start: int, char_end: int) -> list[int]:
        pages = set()
        for start_off, end_off, page_num in page_offsets:
            if char_start < end_off and char_end > start_off:
                pages.add(page_num)
        return sorted(pages)

    for unit_text, unit_start, unit_end in natural_units:
        unit_len = len(unit_text)
        if current_len + unit_len <= chunk_size:
            current_chunk_units.append((unit_text, unit_start, unit_end))
            current_len += unit_len + 1
        else:
            # Finish current chunk
            if current_chunk_units:
                chunk_text = " ".join(u[0] for u in current_chunk_units)
                chunk_start = current_chunk_units[0][1]
                chunk_end = current_chunk_units[-1][2]
                all_chunks.append((chunk_text, chunk_start, chunk_end))

                # Build overlap, but don't duplicate an entire oversized unit
                overlap_units = []
                overlap_len = 0
                for u in reversed(current_chunk_units):
                    if len(u[0]) > overlap_chars and not overlap_units:
                        # This single unit is already larger than overlap_chars;
                        # start fresh without it to avoid full duplication
                        break
                    overlap_units.insert(0, u)
                    overlap_len += len(u[0]) + 1
                    if overlap_len >= overlap_chars:
                        break
                current_chunk_units = overlap_units
                current_len = overlap_len

            # Add the new unit to the current (possibly fresh) chunk
            current_chunk_units.append((unit_text, unit_start, unit_end))
            current_len += unit_len + 1

    # Last chunk
    if current_chunk_units:
        chunk_text = " ".join(u[0] for u in current_chunk_units)
        chunk_start = current_chunk_units[0][1]
        chunk_end = current_chunk_units[-1][2]
        all_chunks.append((chunk_text, chunk_start, chunk_end))

    # Convert to Document objects
    for chunk_text, start_char, end_char in all_chunks:
        pages = _get_pages(start_char, end_char)
        if pages:
            documents.append(
                Document(page_content=chunk_text.strip(), metadata={"page": pages})
            )

    documents.extend(table_chunks)

    logger.info(
        "PDF parsed with pdfplumber: %d text chunks + %d table chunks",
        len(documents) - len(table_chunks),
        len(table_chunks),
    )
    return documents


def _parse_docx_documents(file_bytes: bytes, filename: str) -> list[Document]:
    """Extract DOCX text and split into ~500-char chunks with 50-char overlap."""
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        full_text = _clean_text(" ".join(paragraphs))
        if not full_text:
            return []
        if len(full_text) < 500:
            return [Document(page_content=full_text, metadata={})]

        chunk_size = 500
        overlap = 50
        documents: list[Document] = []
        start = 0
        while start < len(full_text):
            chunk = full_text[start : start + chunk_size].strip()
            if chunk:
                documents.append(Document(page_content=chunk, metadata={}))
            start += chunk_size - overlap
            if start >= len(full_text):
                break
        return documents
    except Exception as exc:
        logger.error("Failed to parse DOCX file %s: %s", filename, exc, exc_info=True)
        return []


def _parse_tabular_documents(file_bytes: bytes, ext: str, filename: str) -> list[Document]:
    """Parse CSV, Excel, or JSON into one document string per row."""
    rows: list[Any] = []
    if ext == "json":
        try:
            raw = json.loads(file_bytes.decode("utf-8-sig"))
            if isinstance(raw, list):
                rows = raw
            elif isinstance(raw, dict):
                rows = raw.get("data") or raw.get("records") or raw.get("rows") or [raw]
            else:
                rows = [raw]
        except Exception as exc:
            logger.error("Failed to parse JSON file %s: %s", filename, exc, exc_info=True)
            return []
    elif ext == "csv":
        try:
            df = pd.read_csv(io.BytesIO(file_bytes))
            rows = df.to_dict(orient="records")
        except Exception as exc:
            logger.error("Failed to parse CSV file %s: %s", filename, exc, exc_info=True)
            return []
    elif ext in ("xlsx", "xls"):
        try:
            df = pd.read_excel(io.BytesIO(file_bytes))
            rows = df.to_dict(orient="records")
        except Exception as exc:
            logger.error("Failed to parse Excel file %s: %s", filename, exc, exc_info=True)
            return []
    else:
        return []

    documents: list[Document] = []
    for row in rows:
        text = _row_to_document(row)
        if text:
            documents.append(Document(page_content=text, metadata={}))
    return documents


def parse_upload_to_documents(file_bytes: bytes, filename: str) -> list[Document]:
    """Parse CSV, Excel, JSON, PDF, or DOCX upload into LangChain Document objects."""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    logger.info("Processing admin RAG file: filename=%s, extension=%s", filename, ext)
    if ext in ("json", "csv", "xlsx", "xls"):
        return _parse_tabular_documents(file_bytes, ext, filename)
    if ext == "pdf":
        return _parse_pdf_documents(file_bytes, filename)
    if ext == "docx":
        return _parse_docx_documents(file_bytes, filename)
    logger.warning("Unsupported admin RAG file type: %s", filename)
    return []


def ingest_documents(
    file_bytes: bytes,
    filename: str,
    mode: Literal["replace", "append"] = "replace",
) -> dict[str, Any]:
    """
    Parse an uploaded file and build/update dense (FAISS) + sparse (BM25) indices.

    - `replace`: wipe existing data before ingest (default).
    - `append`: add new documents to the existing index.
    """
    global _vector_store, _documents, _bm25

    if len(file_bytes) > MAX_UPLOAD_BYTES:
        raise ValueError(f"File exceeds {MAX_UPLOAD_BYTES // (1024 * 1024)} MB limit.")

    documents = parse_upload_to_documents(file_bytes, filename)
    if not documents:
        raise ValueError("No rows found in the uploaded file.")

    t0 = time.perf_counter()
    try:
        with _lock:
            if mode == "replace":
                if os.path.isdir(PERSIST_DIR):
                    shutil.rmtree(PERSIST_DIR, ignore_errors=True)
                os.makedirs(PERSIST_DIR, exist_ok=True)

                embeddings = _get_embeddings()
                vector_store = FAISS.from_texts(["init"], embeddings)
                vector_store.add_documents(documents)
                vector_store.save_local(PERSIST_DIR)
                _vector_store = vector_store

                _documents = list(documents)
                tokenized = [_tokenize(doc.page_content) for doc in _documents]
                _bm25 = BM25Okapi(tokenized)

            else:  # append
                if not _documents:
                    _load_bm25_from_disk()
                store = _get_admin_vector_store()
                if store is None:
                    os.makedirs(PERSIST_DIR, exist_ok=True)
                    store = FAISS.from_texts(["init"], _get_embeddings())
                    store.add_documents(documents)
                    store.save_local(PERSIST_DIR)
                    _vector_store = store
                else:
                    store.add_documents(documents)
                    store.save_local(PERSIST_DIR)
                    _vector_store = store

                _documents.extend(documents)
                tokenized = [_tokenize(doc.page_content) for doc in _documents]
                _bm25 = BM25Okapi(tokenized)

            _save_bm25_to_disk()

    except Exception:
        logger.error(
            "Admin RAG FAISS indexing failed for %s:\n%s",
            filename,
            traceback.format_exc(),
        )
        raise

    elapsed_ms = (time.perf_counter() - t0) * 1000
    logger.info(
        "Admin RAG ingest (mode=%s): %d documents indexed in %.0fms (total=%d)",
        mode,
        len(documents),
        elapsed_ms,
        len(_documents),
    )
    return {
        "status": "success",
        "documents_indexed": len(documents),
        "total_documents": len(_documents),
        "mode": mode,
        "elapsed_ms": round(elapsed_ms),
    }


def _reciprocal_rank_fusion(rankings: list[list[int]], top_k: int) -> list[int]:
    scores: dict[int, float] = {}
    for ranking in rankings:
        for rank, doc_idx in enumerate(ranking):
            scores[doc_idx] = scores.get(doc_idx, 0.0) + 1.0 / (RRF_K + rank + 1)
    ordered = sorted(scores.items(), key=lambda item: item[1], reverse=True)
    return [idx for idx, _ in ordered[:top_k]]


def hybrid_retrieve(query: str, top_k: int = 10) -> list[Document]:
    """
    Hybrid search: top dense (vector) + top sparse (BM25) candidates merged via RRF.
    Returns the top_k Document objects (page_content + metadata, e.g. page number).
    """
    if not _documents:
        _load_bm25_from_disk()

    if not _documents or _bm25 is None:
        logger.warning("Admin RAG hybrid_retrieve: no indexed documents")
        return []

    rankings: list[list[int]] = []

    # Sparse (BM25)
    query_tokens = _tokenize(query)
    if query_tokens:
        bm25_scores = _bm25.get_scores(query_tokens)
        sparse_indices = sorted(
            range(len(bm25_scores)),
            key=lambda i: bm25_scores[i],
            reverse=True,
        )
        sparse_top = [
            i for i in sparse_indices[:SPARSE_CANDIDATES] if bm25_scores[i] > 0
        ]
        if sparse_top:
            rankings.append(sparse_top)

    # Dense (vector)
    store = _get_admin_vector_store()
    if store is not None:
        dense_docs = store.similarity_search(query, k=DENSE_CANDIDATES)
        dense_indices: list[int] = []
        for doc in dense_docs:
            try:
                dense_indices.append(
                    next(
                        i
                        for i, stored in enumerate(_documents)
                        if stored.page_content == doc.page_content
                    )
                )
            except StopIteration:
                continue
        if dense_indices:
            rankings.append(dense_indices)

    if not rankings:
        return []

    merged_indices = _reciprocal_rank_fusion(rankings, top_k=top_k)
    final_chunks = [
        _documents[i] for i in merged_indices if 0 <= i < len(_documents)
    ]

    logger.info(
        f"RAG final chunks (top {top_k}): pages={[doc.metadata.get('page', '?') for doc in final_chunks]}"
    )

    logger.info(
        "Admin RAG hybrid_retrieve: query=%r dense=%d sparse=%d merged=%d",
        query[:80],
        len(rankings[1]) if len(rankings) > 1 else 0,
        len(rankings[0]) if rankings else 0,
        len(final_chunks),
    )
    for i, doc in enumerate(final_chunks):
        logger.info(
            "RAG chunk %d: page=%s, content=%s...",
            i + 1,
            doc.metadata.get("page", "N/A"),
            doc.page_content[:100],
        )
    return final_chunks


def get_rag_index_status() -> dict[str, Any]:
    """Return current admin RAG index status for the dashboard."""
    if not _documents:
        _load_bm25_from_disk()
    return {
        "documents_indexed": len(_documents),
        "collection": "faiss",
        "vector_store_ready": _get_admin_vector_store() is not None,
        "bm25_ready": _bm25 is not None and len(_documents) > 0,
    }


def warmup_admin_rag() -> None:
    """Load persisted BM25 index at startup (vector store loads lazily)."""
    _load_bm25_from_disk()
    if _documents:
        _get_admin_vector_store()